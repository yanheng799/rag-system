"""Word (.docx) 文档解析器（基于 python-docx）"""

from __future__ import annotations

import logging

from docx import Document

from src.ingestion.parsers.base import BaseParser, ParsedElement, ParseError

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """使用 python-docx 解析 Word 文档"""

    def __init__(self, extract_images: bool = True):
        self._do_extract_images = extract_images

    def parse(self, file_path: str) -> list[ParsedElement]:
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ParseError(file_path, str(e)) from e

        elements: list[ParsedElement] = []
        position = 0  # 文档中的顺序位置

        # 预加载图片关系映射
        image_rels = {}
        if self._do_extract_images:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        blob = rel.target_part.blob
                        ct = rel.target_part.content_type
                        ext = self._content_type_to_ext(ct)
                        image_rels[rel.rId] = {"blob": blob, "ext": ext}
                    except Exception:
                        continue

        img_index = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # 段落处理
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break

                if para:
                    # TOC 条目：跳过段落文本及其中的图片
                    para_type = self._detect_paragraph_type(para)
                    if para_type is None:
                        continue

                    # 检测段落中的图片（按 rId 去重，同一图片只提取一次）
                    if self._do_extract_images:
                        seen_rids: set[str] = set()
                        blips = element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                        for blip in blips:
                            rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            if rid and rid not in seen_rids and rid in image_rels:
                                seen_rids.add(rid)
                                img_info = image_rels[rid]
                                filename = f"img_{img_index}.{img_info['ext']}"
                                elements.append(
                                    ParsedElement(
                                        elem_type="image",
                                        content=f"[图片: {filename}]",
                                        page=1,
                                        bbox=(0, position, 0, position + 1),
                                        style={"para_index": position},
                                        raw={
                                            "ext": img_info["ext"],
                                            "filename": filename,
                                            "image_bytes": img_info["blob"],
                                        },
                                    )
                                )
                                img_index += 1
                                position += 1

                    if para.text.strip():
                        text = para.text.strip()
                        elements.append(
                            ParsedElement(
                                elem_type=para_type,
                                content=text,
                                page=1,
                                bbox=(0, position, 0, position + 1),
                                style=self._extract_paragraph_style(para),
                            )
                        )
                        position += 1

            elif tag == "tbl":
                # 表格处理
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break

                if table:
                    table_text = self._extract_table_text(table)
                    if table_text.strip():
                        elements.append(
                            ParsedElement(
                                elem_type="table",
                                content=table_text,
                                page=1,
                                bbox=(0, position, 0, position + 1),
                                style={"table_index": position},
                                raw=table,
                            )
                        )
                        position += 1

        logger.info("Word 解析完成: %s, 共 %d 个元素", file_path, len(elements))
        return elements

    def _detect_paragraph_type(self, para) -> str | None:
        """判断段落类型，返回 None 表示应跳过（如 TOC 条目）"""
        style_name = para.style.name if para.style else ""
        if "toc" in style_name.lower() or "目录" in style_name:
            return None
        if "Heading" in style_name or "标题" in style_name:
            return "title"
        if "List" in style_name or "列表" in style_name:
            return "list_item"
        return "text"

    def _extract_paragraph_style(self, para) -> dict:
        """提取段落样式信息"""
        style = {}
        if para.style:
            style["style_name"] = para.style.name
        if para.runs:
            run = para.runs[0]
            style["bold"] = run.bold or False
            style["font_size"] = run.font.size.pt if run.font.size else None
        return style

    def _extract_table_text(self, table) -> str:
        """提取表格内容为 Markdown 表格"""
        md_lines = []
        for i, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                text = cell.text.strip().replace("|", "｜").replace("\n", "<br>") if cell.text else ""
                cells.append(text)
            md_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                md_lines.append("|" + "|".join("---" for _ in cells) + "|")
        return "\n".join(md_lines)

    @staticmethod
    def _content_type_to_ext(content_type: str) -> str:
        mapping = {
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
        }
        return mapping.get(content_type, "png")

    def supported_types(self) -> list[str]:
        return ["docx"]
