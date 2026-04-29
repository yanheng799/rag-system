"""Word (.docx) 文档解析器（基于 python-docx）"""

from __future__ import annotations

import logging

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ingestion.parsers.base import BaseParser, ParsedElement, ParseError

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """使用 python-docx 解析 Word 文档"""

    def parse(self, file_path: str) -> list[ParsedElement]:
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ParseError(file_path, str(e))

        elements: list[ParsedElement] = []
        position = 0  # 文档中的顺序位置

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # 段落处理
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break

                if para and para.text.strip():
                    text = para.text.strip()
                    elem_type = self._detect_paragraph_type(para)
                    elements.append(
                        ParsedElement(
                            elem_type=elem_type,
                            content=text,
                            page=1,  # Word 无精确页码概念
                            bbox=(0, position, 0, position + 1),
                            style=self._extract_paragraph_style(para),
                        )
                    )
                    position += 1

            elif tag == "tbl":
                # 表格处理
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break

                if table:
                    table_text = self._extract_table_text(table)
                    if table_text.strip():
                        elements.append(
                            ParsedElement(
                                elem_type="table",
                                content=table_text,
                                page=1,
                                bbox=(0, position, 0, position + 1),
                                style={"table_index": position},
                                raw=table,
                            )
                        )
                        position += 1

        logger.info("Word 解析完成: %s, 共 %d 个元素", file_path, len(elements))
        return elements

    def _detect_paragraph_type(self, para) -> str:
        """判断段落类型"""
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name or "标题" in style_name:
            return "title"
        if "List" in style_name or "列表" in style_name:
            return "list_item"
        return "text"

    def _extract_paragraph_style(self, para) -> dict:
        """提取段落样式信息"""
        style = {}
        if para.style:
            style["style_name"] = para.style.name
        if para.runs:
            run = para.runs[0]
            style["bold"] = run.bold or False
            style["font_size"] = run.font.size.pt if run.font.size else None
        return style

    def _extract_table_text(self, table) -> str:
        """提取表格内容为 Markdown 表格"""
        md_lines = []
        for i, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                text = cell.text.strip().replace("|", "｜") if cell.text else ""
                cells.append(text)
            md_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                md_lines.append("|" + "|".join("---" for _ in cells) + "|")
        return "\n".join(md_lines)

    def supported_types(self) -> list[str]:
        return ["docx"]
